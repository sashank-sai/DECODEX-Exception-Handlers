"""
DECODE X 2026 — Hackathon Submission Document Builder v2
=========================================================
Enhanced: Detailed methodology behind each chart + highlighted key findings.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from pathlib import Path

DATA_DIR = Path(r'c:\Users\asus\Desktop\decodex')
CHART_DIR = DATA_DIR / 'charts'
OUT_FILE = DATA_DIR / 'DECODEX_2026_Submission_v2.docx'

doc = Document()

# ============================================================
# STYLES
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x26, 0x32, 0x38)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)
    if level == 1: hs.font.size = Pt(22); hs.font.bold = True
    elif level == 2: hs.font.size = Pt(16); hs.font.bold = True
    elif level == 3: hs.font.size = Pt(13); hs.font.bold = True

def add_line(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single', qn('w:sz'): '6', qn('w:space'): '1', qn('w:color'): '0D47A1'})
    pBdr.append(bottom); pPr.append(pBdr)

def add_chart(doc, filename, caption='', width=Inches(6.2)):
    path = CHART_DIR / filename
    if path.exists():
        doc.add_picture(str(path), width=width)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            cap = doc.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.runs[0].italic = True; cap.runs[0].font.size = Pt(9)
            cap.runs[0].font.color.rgb = RGBColor(0x61, 0x61, 0x61)

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Shading Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]; cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold = True; r.font.size = Pt(10)
    for ri, rd in enumerate(rows):
        for ci, cv in enumerate(rd):
            cell = table.rows[ri+1].cells[ci]; cell.text = str(cv)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs: r.font.size = Pt(10)
    doc.add_paragraph()

def add_highlight_box(doc, title, points):
    """Add a highlighted key findings box with bold heading + bullet points."""
    p = doc.add_paragraph()
    run = p.add_run(f'🔑 {title}')
    run.bold = True; run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0xD3, 0x2F, 0x2F)
    for point in points:
        bp = doc.add_paragraph(style='List Bullet')
        # Split on ** markers for bold emphasis
        parts = point.split('**')
        for i, part in enumerate(parts):
            r = bp.add_run(part)
            r.font.size = Pt(11)
            if i % 2 == 1:  # odd parts are between ** markers = bold
                r.bold = True
                r.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)

def add_methodology_box(doc, text):
    """Add a methodology explanation in a distinct gray italic block."""
    p = doc.add_paragraph()
    run = p.add_run('📊 How this chart was built: ')
    run.bold = True; run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x42, 0x42, 0x42)
    run2 = p.add_run(text)
    run2.font.size = Pt(10); run2.italic = True
    run2.font.color.rgb = RGBColor(0x61, 0x61, 0x61)

def add_bold_para(doc, text):
    """Add paragraph with **bold** markers parsed."""
    p = doc.add_paragraph()
    parts = text.split('**')
    for i, part in enumerate(parts):
        r = p.add_run(part)
        r.font.size = Pt(11)
        if i % 2 == 1: r.bold = True; r.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)
    return p

# ============================================================
# COVER PAGE
# ============================================================
for _ in range(4): doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('DECODE X 2026')
run.font.size = Pt(36); run.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1); run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Dubai RTA Bus Network\nStrategic Advisory Submission')
run.font.size = Pt(20); run.font.color.rgb = RGBColor(0x42, 0x42, 0x42)

doc.add_paragraph()
tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tagline.add_run('Mobility Shift: Demand Diagnostics, Forecasting & Fleet Optimization')
run.font.size = Pt(13); run.font.color.rgb = RGBColor(0xF5, 0x7C, 0x00); run.italic = True

doc.add_paragraph(); doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run(
    'Stage 1 — Pre-Shock Analysis\n'
    'Advisory Mandate: Multi-Year Demand Diagnostics & Route-Level Forecasting\n\n'
    'Data Scope: 12 Routes | 56 Stops | 7 Zones | 1,277 Days\n'
    'Dataset: 195,381 route-stop-day observations (Jan 2022 – Jun 2025)')
run.font.size = Pt(11); run.font.color.rgb = RGBColor(0x61, 0x61, 0x61)

doc.add_page_break()

# ============================================================
# EXECUTIVE SUMMARY
# ============================================================
doc.add_heading('Executive Summary', level=1)
add_line(doc)

add_bold_para(doc,
    'This report delivers a **complete demand diagnostic** of Dubai RTA\'s bus network '
    'spanning 2022–H1 2025, a **route-level daily forecast** for H2 2025, and '
    '**budget-neutral fleet optimization** proposals. The analysis covers **12 routes** '
    'across **7 zones** serving a projected **11.1 million annual passengers**.')

add_table(doc,
    ['Metric', 'Value', 'Significance'],
    [
        ['Annual Demand (2024)', '10.3M passengers', 'All-time high'],
        ['Growth Rate', '+13.1% (2023→24)', 'Accelerating year-over-year'],
        ['2025 Full-Year Projection', '11.1M passengers', '+7.9% vs 2024'],
        ['Seasonal Swing', '31% winter-to-summer', 'Strongest demand driver'],
        ['Most Overloaded Route', 'X28 (Express)', 'Risk score: 89/100'],
        ['Most Underutilized Route', 'F18 (Feeder)', 'Waste score: 88/100'],
        ['Reallocation Opportunity', '11 buses moved', 'Zero new capital required'],
        ['Top Bottleneck', 'Stop 31 (Coastal Marina)', '326 pax per dwell-minute'],
    ])

add_highlight_box(doc, 'Bottom Line', [
    'Express corridors are **systemically overloaded** — all 3 routes exceed warning thresholds.',
    'A **budget-neutral redistribution of 11 buses** can reduce X28 headway by **38%** immediately.',
    'Without intervention, **all 12 routes** will exceed their 95th percentile demand during Winter 2025.',
])

doc.add_page_break()

# ============================================================
# SECTION 1: MULTI-YEAR DEMAND DYNAMICS
# ============================================================
doc.add_heading('1. Diagnose Multi-Year Demand Dynamics (2022–H1 2025)', level=1)
add_line(doc)

# --- Chart: Yearly Growth ---
doc.add_heading('1.1 Year-over-Year Growth Trajectory', level=2)

add_methodology_box(doc,
    'We aggregated Total_Pax (Boarding_Count + Alighting_Count) by calendar year across all 12 routes '
    'and 56 stops. Each bar represents the cumulative annual ridership. Growth percentages are computed '
    'as (Year_N / Year_N-1 - 1) × 100. The 2025 bar shows only H1 actual data (Jan–Jun). '
    'This simple year-over-year comparison reveals whether growth is linear, accelerating, or decelerating.')

add_chart(doc, '02_yearly_growth.png',
    'Figure 1: Yearly passenger totals with YoY growth rates — 2025 shows H1 only')

add_bold_para(doc,
    'The network served **8.4M passengers in 2022**, growing to **9.2M in 2023** (+9.2%) '
    'and **10.3M in 2024** (+13.1%). The growth rate is not constant — it is **accelerating**. '
    '2025 H1 data (5.6M) annualizes to ~11.1M, suggesting **+7.9% full-year growth**.')

add_highlight_box(doc, 'Key Findings', [
    'Growth rate **accelerated from +9.2% to +13.1%** in one year — this is not linear growth.',
    'The network added approximately **2 million annual riders** in just two years.',
    'Even the conservative 2025 projection (+7.9%) represents an additional **820,000 passengers** over 2024.',
    'At this trajectory, the system will serve **~13M passengers by 2027** without capacity additions.',
])

# --- Chart: Growth Decomposition ---
doc.add_heading('1.2 Growth Decomposition by Route Type', level=2)

add_methodology_box(doc,
    'We decomposed the yearly total into four route-type segments (City, Express, Feeder, Intercity) '
    'using a stacked bar chart. For 2025, we combined H1 actual data with H2 forecast predictions '
    '(hatched portion indicates forecast). The stacking reveals which segments are driving growth '
    'and whether growth is broad-based or concentrated in a single category.')

add_chart(doc, '13_growth_decomposition.png',
    'Figure 2: Growth decomposition by route type — hatched area = H2 2025 forecast')

add_bold_para(doc,
    '**All four route types are growing in parallel** — no segment is shrinking or stagnant. '
    'City routes dominate total volume (growing from 3.6M to projected 5.5M), while Express '
    'routes show the **steepest per-route growth rate** because they serve high-density corridors '
    'with fewer routes.')

add_highlight_box(doc, 'Key Findings', [
    '**City routes** are the backbone: 4 routes carrying 5.5M projected passengers (49% of network).',
    '**Express routes** punch above their weight: 3 routes but highest per-stop intensity.',
    'Growth is **structural and broad-based** — not driven by a single event, route type, or season.',
    '**Intercity** is the smallest contributor (~1.8M), suggesting potential for capacity optimization.',
])

# --- Chart: Monthly Timeline ---
doc.add_heading('1.3 Monthly Demand Timeline with Forecast', level=2)

add_methodology_box(doc,
    'Monthly passenger totals are computed by summing Total_Pax across all route-stop combinations per month. '
    'The blue line represents actual historical data (Jan 2022 – Jun 2025). The red dashed line represents '
    'our H2 2025 forecast (Jul – Dec), generated using a component decomposition model that combines '
    'linear growth trend, seasonal multipliers, day-of-week effects, and congestion elasticity. '
    'Blue/red shaded bands indicate Winter Peak and Summer Low seasons respectively.')

add_chart(doc, '01_monthly_demand_trend.png',
    'Figure 3: Monthly demand — blue = actual, red dashed = H2 2025 forecast, shaded = seasonal zones')

add_highlight_box(doc, 'Key Findings', [
    '**Cyclical pattern repeats every year**: peaks in Nov–Mar (blue shading), troughs in Jun–Aug (red).',
    'Each successive year\'s peak is **higher than the previous** — the amplitude is growing.',
    '**January 2025 set the all-time monthly record** at 1.07M passengers.',
    'The forecast predicts **December 2025 will reach ~1.09M**, breaking the January record.',
])

doc.add_page_break()

# ============================================================
# SECTION 2: SEASONALITY & WEEKDAY/WEEKEND
# ============================================================
doc.add_heading('2. Detect Seasonality & Weekday/Weekend Divergence', level=1)
add_line(doc)

# --- Chart: Seasonal Pattern ---
doc.add_heading('2.1 Seasonal Demand Comparison', level=2)

add_methodology_box(doc,
    'Each day in the dataset was classified into one of three Dubai-specific seasons: '
    'Winter Peak (Nov–Mar, pleasant weather + tourism), Summer Moderate (Jun–Aug, extreme heat), '
    'and Shoulder (Apr–May, Sep–Oct, transition months). We then calculated the average '
    'Total_Pax per stop-day for each season. The bar chart compares these three averages to '
    'quantify the seasonal swing that RTA must plan for.')

add_chart(doc, '03_seasonal_demand.png',
    'Figure 4: Seasonal demand pattern — Winter is 31% above Summer')

add_bold_para(doc,
    'Winter Peak averages **189.5 passengers per stop-day**, Shoulder averages **166.6**, '
    'and Summer Moderate averages **145.1**. The winter-to-summer gap of **31%** is the single '
    'largest predictable demand driver in the network — larger than any route or zone effect.')

add_highlight_box(doc, 'Key Findings', [
    'A **31% seasonal swing** means the network needs fundamentally different service levels by season.',
    'Winter Peak (5 months) contributes approximately **45% of annual demand**.',
    'Summer\'s lower demand creates a natural window for **fleet maintenance and driver training**.',
    'This pattern aligns with Dubai\'s tourism calendar — winter events (shopping festivals, Expo) amplify it.',
])

# --- Chart: Season Decomposition ---
add_chart(doc, '14_growth_decomposition_season.png',
    'Figure 5: Growth decomposition by season — all three seasons grow year-over-year')

add_methodology_box(doc,
    'Identical stacking methodology as Figure 2, but segmented by season instead of route type. '
    'This shows whether growth is concentrated in one season or spread across all three. '
    'For 2025, H1 actual + H2 forecast are combined; hatched = forecast portion.')

add_highlight_box(doc, 'Key Findings', [
    'All three seasons grow each year — growth is **not just a tourism/winter phenomenon**.',
    'The **winter-summer gap is widening** each year, amplifying the seasonal resource challenge.',
    'Summer demand, while lowest, is **still growing** — baseline ridership is structurally rising.',
])

# --- Chart: Weekday vs Weekend ---
doc.add_heading('2.2 Weekday vs Weekend Divergence', level=2)

add_methodology_box(doc,
    'Dubai operates on a Friday–Saturday weekend (Islamic calendar). We classified each day as '
    'Weekday (Sun–Thu) or Weekend (Fri–Sat) and computed the average per-stop-day ridership '
    'for each route. The grouped bar chart shows both values side-by-side — the gap reveals '
    'how dependent each route is on commuter (work) travel vs discretionary travel.')

add_chart(doc, '07_weekday_weekend.png',
    'Figure 6: Weekday vs Weekend demand by route — Express routes show steepest drop')

add_bold_para(doc,
    'Every route shows **higher weekday demand** than weekend demand —  this is a **commuter-dominated network**. '
    'Express routes show the **steepest weekend drop (~8%)**, while City routes are more stable (3–5%). '
    'Color-coded x-axis labels match route types: blue = City, red = Express, green = Feeder, orange = Intercity.')

add_highlight_box(doc, 'Key Findings', [
    '**Express routes are almost purely commuter** — 8% demand drop on weekends.',
    '**City routes serve mixed purposes** (commute + errands + tourism) with only 3-5% weekend drop.',
    'Weekend service on Express routes can be **safely reduced by 7-8%** and buses redeployed to weekday peaks.',
])

# --- Chart: DOW Heatmap ---
doc.add_heading('2.3 Day-of-Week Granularity', level=2)

add_methodology_box(doc,
    'We computed the average daily total demand by route × day-of-week for H1 2025 data (the most recent '
    '6-month period). The result is pivoted into a Route × DayOfWeek matrix and rendered as a heatmap '
    'where darker red = higher demand. This reveals which specific days carry the most load for each route, '
    'enabling granular timetabling rather than a simple weekday/weekend binary split.')

add_chart(doc, '12_dow_heatmap.png',
    'Figure 7: Demand heatmap — darker = higher demand — Wednesday is peak, Saturday is lowest')

add_highlight_box(doc, 'Key Findings', [
    '**Wednesday is consistently the peak day** across most routes — mid-week office presence peak.',
    '**Saturday is the lowest** demand day network-wide, followed by Friday.',
    'Timetabling should use **3 tiers** (not binary): Peak (Sun–Thu), Reduced (Fri), Minimal (Sat).',
    'Express routes show the **widest color range** (most variable) vs City routes (most stable).',
])

doc.add_page_break()

# ============================================================
# SECTION 3: CORRIDOR STRUCTURAL IMBALANCE
# ============================================================
doc.add_heading('3. Identify Corridor-Level Structural Imbalance', level=1)
add_line(doc)

# --- Chart: Route Type Comparison ---
doc.add_heading('3.1 Route Type Performance Comparison', level=2)

add_methodology_box(doc,
    'Three metrics are computed per route type: (1) Total passenger volume = sum of Total_Pax across '
    'all years, (2) Average per-stop-day = mean Total_Pax per route-stop-day observation, and '
    '(3) Number of routes in each category. These three perspectives reveal whether resource allocation '
    '(number of routes) matches demand distribution — a core structural imbalance diagnostic.')

add_chart(doc, '04_route_type_comparison.png',
    'Figure 8: Three-panel route type analysis — volume, intensity, and count')

add_highlight_box(doc, 'Key Findings', [
    '**Express routes have the highest per-stop intensity** (197 avg) despite having only 3 routes.',
    '**City routes carry the most total volume** (12.9M) with 4 routes — they are the network backbone.',
    '**Intercity routes have the lowest metrics** across both volume and intensity — potential overallocation.',
    'The imbalance is clear: **Express is under-provisioned, Intercity is over-provisioned**.',
])

# --- Chart: Overload Risk ---
doc.add_heading('3.2 Overload Risk Scoring', level=2)

add_methodology_box(doc,
    'For each route, we calculated Passengers-per-Kilometer = (Average Daily Total Pax) / (Route Length in km) '
    'using H1 2025 data. This was then normalized to a 0–100 scale where 0 = lowest density route and '
    '100 = highest density route. Two threshold lines are drawn: CRITICAL (>70) indicating immediate intervention '
    'needed, and WARNING (>40) indicating monitoring required. Colors match route types for pattern visibility.')

add_chart(doc, '09_overload_risk.png',
    'Figure 9: Overload risk scores — red threshold lines at CRITICAL (70) and WARNING (40)')

add_bold_para(doc,
    'Route X28 (Express) scores **89/100** — the highest overload risk in the network, carrying '
    '**18.8 passengers per kilometer**. All three Express routes are above the WARNING line. '
    'This is not a single-route problem — it is a **systemic Express category failure**.')

add_highlight_box(doc, 'Key Findings', [
    '**ALL 3 Express routes exceed the WARNING threshold** — this is a category-wide structural issue.',
    '**X28 is in CRITICAL territory** (89/100 risk) with 18.8 pax/km — the densest corridor.',
    'Feeder and Intercity routes score below 30 — they have **headroom to donate capacity**.',
    'The gap between Express (60-89) and others (<30) represents **misallocated fleet resources**.',
])

# --- Overloaded corridors table ---
doc.add_heading('Overloaded Corridors (Top 3)', level=3)
add_table(doc,
    ['Route', 'Type', 'Risk Score', 'Pax/km', 'Forecast Growth', 'Priority'],
    [
        ['X28', 'Express', '89/100', '18.8', '+8.0%', 'CRITICAL'],
        ['X11', 'Express', '68/100', '10.3', '+8.4%', 'HIGH'],
        ['X66', 'Express', '60/100', '12.3', '+7.2%', 'HIGH'],
    ])

# --- Underutilized corridors table ---
doc.add_heading('Underutilized Corridors (Top 3)', level=3)
add_table(doc,
    ['Route', 'Type', 'Waste Score', 'Pax/km', 'Length (km)', 'Stops'],
    [
        ['F18', 'Feeder', '88/100', '10.0', '18.1', '6'],
        ['E16', 'Intercity', '75/100', '7.6', '22.1', '5'],
        ['F12', 'Feeder', '75/100', '6.7', '27.0', '5'],
    ])

# --- Chart: Zone Demand ---
doc.add_heading('3.3 Zone Demand Distribution', level=2)

add_methodology_box(doc,
    'Total and average passengers are aggregated by the Zone field from Bus_Stops.csv — each stop '
    'belongs to one of 7 zones representing geographic/functional areas of Dubai. Bars are sorted by '
    'total volume and color-mapped on a warm gradient (yellow→red) for visual intensity. The "(avg: N)" '
    'annotation reveals per-stop pressure — a zone can rank low in total but high in intensity if it '
    'has very few stops.')

add_chart(doc, '06_zone_demand.png',
    'Figure 10: Zone demand — Downtown leads in total, Coastal Marina leads in per-stop intensity')

add_highlight_box(doc, 'Key Findings', [
    '**CBD Downtown** leads in total volume (8.7M) — the primary trip origin/destination.',
    '**Coastal Marina** has the **highest per-stop intensity (199 avg)** with only 5 stops — a capacity crunch.',
    '**Jebel Ali** (industrial) has the lowest demand — resources here may be partially redirectable.',
])

# --- Chart: Top Stops ---
doc.add_heading('3.4 Stop-Level Bottleneck Identification', level=2)

add_methodology_box(doc,
    'Total passengers were aggregated by Stop_ID across all dates and routes. The top 10 stops are '
    'displayed as a horizontal bar chart color-coded by zone. The stop type (Regular, Interchange, '
    'Metro_Link, Terminal) is noted — this reveals whether bottlenecks occur at designed interchange '
    'points or at ordinary mid-route stops (which would indicate organic demand outpacing infrastructure).')

add_chart(doc, '08_top_stops.png',
    'Figure 11: Top 10 busiest stops — color = zone, label = stop type')

add_bold_para(doc,
    'Stop 52 (Coastal Marina, **Regular** stop) is the single busiest in the network with **1.26M total passengers** '
    '— it outperforms Metro Link interchange stations. This indicates that **organic demand at a regular stop** '
    'has exceeded the capacity originally designed for that stop type.')

add_highlight_box(doc, 'Key Findings', [
    '**3 of the top 10 stops** are in Coastal Marina — a zone with only 5 stops total.',
    '**Regular stops outperform Interchange stations** — demand has shifted to non-designed hotspots.',
    '**Metro Link Stop 20** (1.26M) is the #2 stop — bus-metro integration is a critical interchange.',
    'Stop 31 reaches **326 pax per dwell-minute** — the single most congested boarding point.',
])

doc.add_page_break()

# ============================================================
# SECTION 4: CONGESTION-DEMAND INTERACTION
# ============================================================
doc.add_heading('4. Model Congestion–Demand Interaction', level=1)
add_line(doc)

add_methodology_box(doc,
    'We joined the Train_Traffic dataset (daily Congestion_Level, Avg_Speed_kmph) with the ridership data. '
    'Each dot in the scatter plot represents one day. X-axis = Congestion Level (1–5), Y-axis = total network '
    'daily passengers. The color gradient (green→red) encodes average bus speed that day. A linear trend line '
    'is fitted via numpy polyfit to quantify the direction and strength of the congestion-ridership relationship.')

add_chart(doc, '05_congestion_paradox.png',
    'Figure 12: Congestion Paradox — ridership INCREASES with congestion (color = bus speed)')

add_bold_para(doc,
    'This reveals the **"Congestion Paradox"**: when road congestion increases, bus ridership goes **up**, '
    'not down. This is a **modal shift effect** — private car drivers switch to public transit when '
    'congestion makes driving impractical. However, the buses themselves are also stuck in traffic, '
    'with average speeds dropping from **~40 km/h (Level 1) to ~21 km/h (Level 5)**.')

add_highlight_box(doc, 'Key Findings', [
    'Ridership increases by **~20%** when congestion moves from Level 1 to Level 5.',
    'But bus speeds **halve** in the same range — each trip takes twice as long.',
    'This creates a **vicious cycle**: more passengers + slower buses = overcrowding AND delays.',
    'This is the strongest data-driven argument for **dedicated bus-priority lanes**.',
    'Demand forecasting models **must include congestion as an input** — ignoring it underestimates demand by 15-20%.',
])

doc.add_heading('Strategic Implication', level=2)
add_bold_para(doc,
    'Dedicated bus lanes would **break the vicious cycle** by maintaining speed (capacity) while '
    'absorbing the modal shift ridership. Additionally, **dynamic headway adjustment** should trigger '
    'when congestion exceeds Level 3 — dispatching extra buses to absorb the surge demand '
    'before overcrowding becomes critical.')

doc.add_page_break()

# ============================================================
# SECTION 5: FORECAST
# ============================================================
doc.add_heading('5. Route-Level Daily Demand Forecast: Jul 1 – Dec 31, 2025', level=1)
add_line(doc)

doc.add_heading('5.1 Forecasting Methodology', level=2)

add_bold_para(doc,
    'We use a **component decomposition** approach that models four independent demand drivers, '
    'then combines them **multiplicatively**: Forecast = Base × Trend × Seasonal × DOW × Congestion.')

add_table(doc,
    ['Component', 'Method', 'Key Parameter', 'Why It Matters'],
    [
        ['Growth Trend', 'Linear regression on monthly totals', '+2,100 pax/day/year', 'Captures structural growth'],
        ['Seasonal Effect', 'Monthly multipliers (3-year avg)', 'Winter: 1.12x, Summer: 0.88x', '31% seasonal swing'],
        ['Day-of-Week', 'DOW index per route', 'Wed: 1.06x, Sat: 0.83x', '7-8% weekday/weekend gap'],
        ['Congestion', 'Elasticity coefficient per route', '+4.2% per congestion level', 'Modal shift capture'],
    ])

add_methodology_box(doc,
    'Each component was estimated independently from the 3.5-year historical dataset (2022–H1 2025). '
    'The growth trend was fit using linear regression on monthly aggregates. Seasonal multipliers are '
    '36-month averages of monthly deviation from annual mean. Day-of-week effects are per-route ratios '
    'of each day vs the route\'s weekly average. Congestion elasticity captures the net ridership change '
    'per unit increase in congestion level. Components are combined multiplicatively to generate '
    '2,208 route-day predictions (12 routes × 184 days).')

# --- Chart: Forecast vs Actual ---
doc.add_heading('5.2 Forecast Results', level=2)

add_methodology_box(doc,
    'To validate the forecast magnitude, we compare H2 2025 forecast (red bars, LEFT) against H2 2024 actual '
    '(blue bars, RIGHT) month by month. The forecast bar is placed first for easy reference — you read the '
    'projected demand first, then compare against what actually happened in the same month of the prior year. '
    'Growth percentages above each forecast bar show the expected increase.')

add_chart(doc, '10_forecast_vs_actual.png',
    'Figure 13: H2 2025 Forecast (red, left) vs H2 2024 Actual (blue, right) — consistent +7-8% per month')

add_highlight_box(doc, 'Key Findings', [
    'Forecast total: **5,500,596 passengers** for H2 2025 (+7.5% vs H2 2024).',
    'Full-year 2025 projection: **11,142,733 passengers** (+7.9% vs 2024).',
    'Growth is **uniformly +7-8% every month** — not concentrated in any single month.',
    '**December 2025 is the peak** — projected to be the highest-demand month in network history.',
    '**All 12 routes exceed their historical 95th percentile** during November–December.',
])

add_table(doc,
    ['Metric', 'Value'],
    [
        ['H2 2025 Forecast', '5,500,596 passengers'],
        ['Growth vs H2 2024', '+7.5%'],
        ['Full-Year 2025 Projection', '11,142,733 passengers'],
        ['Peak Month', 'December 2025 (~1.09M pax)'],
        ['Capacity Risk', 'All 12 routes exceed P95 in Nov-Dec'],
    ])

doc.add_page_break()

# ============================================================
# SECTION 5a: RISK IDENTIFICATION
# ============================================================
doc.add_heading('5a. Risk Identification', level=1)
add_line(doc)

doc.add_heading('Overload Risk Corridors', level=2)
add_bold_para(doc,
    'Three Express corridors are flagged as high overload risk using a **composite score** '
    'weighting: Passenger Density (40%) + Per-Stop Intensity (30%) + Forecast Growth (30%). '
    'These routes serve the **highest-demand corridors with the fewest buses**.')

add_table(doc,
    ['Route', 'Type', 'Risk', 'Pax/km', 'Pax/Stop', 'H2 Growth', 'Action Required'],
    [
        ['X28', 'Express', '89', '18.8', '374', '+8.0%', 'Add 3 buses immediately'],
        ['X11', 'Express', '68', '10.3', '437', '+8.4%', 'Contingency bus ready'],
        ['X66', 'Express', '60', '12.3', '330', '+7.2%', 'Review headway schedule'],
    ])

doc.add_heading('Underutilized Capacity', level=2)
add_bold_para(doc,
    'Three routes show **significant capacity waste** — low Pax/km relative to their infrastructure '
    'footprint (long routes, many stops). These are the **donor routes** for reallocation.')

add_table(doc,
    ['Route', 'Type', 'Waste Score', 'Pax/km', 'Length', 'Recommendation'],
    [
        ['F18', 'Feeder', '88', '10.0', '18.1 km', 'Reduce 1 bus'],
        ['E16', 'Intercity', '75', '7.6', '22.1 km', 'Reduce 1 bus'],
        ['F12', 'Feeder', '75', '6.7', '27.0 km', 'Reduce 3 buses'],
    ])

doc.add_heading('Emerging Imbalance: Coastal Marina', level=2)
add_bold_para(doc,
    'The most critical **emerging imbalance** is the Coastal Marina zone. Despite ranking only 5th in total '
    'volume, it has the **highest per-stop intensity** (199 avg pax/stop-day) with only **5 stops**. '
    'Stop 31 reaches **326 pax/dwell-minute** — the single most congested boarding point in the network. '
    'This bottleneck is **worsening** as coastal development drives residential growth in the area.')

doc.add_page_break()

# ============================================================
# SECTION 5b: OPERATIONAL ADJUSTMENTS
# ============================================================
doc.add_heading('5b. Propose Operational Adjustments', level=1)
add_line(doc)

# --- Fleet Reallocation ---
doc.add_heading('Fleet Reallocation (Budget-Neutral)', level=2)

add_methodology_box(doc,
    'Current fleet per route was estimated from peak-hour demand (top 6.25% of daily) divided by bus capacity '
    '(60 passengers). Optimal allocation redistributes the same 81 total buses proportionally to each route\'s '
    'Pax/km density. The grouped bar chart shows Current fleet (blue, left) vs Proposed fleet (colored, right) '
    'with delta badges showing the change. Green = buses added, red = buses removed.')

add_chart(doc, '11_fleet_reallocation.png',
    'Figure 14: Current vs Proposed fleet — blue = current, green/red = proposed, badges = delta')

add_bold_para(doc,
    '**11 buses redistributed, zero new purchases.** The total fleet stays at **81 buses** — '
    'only the distribution changes. The biggest move: X28 goes from **7 to 10 buses**, reducing '
    'headway from **20 minutes to 12.4 minutes — a 38% improvement**.')

add_table(doc,
    ['Route', 'Type', 'Current', 'Proposed', 'Change', 'Headway Impact'],
    [
        ['X28', 'Express', '7', '10', '+3', '20 min → 12.4 min (38% improvement)'],
        ['C03', 'City', '8', '10', '+2', '15 min → 11.5 min (23% improvement)'],
        ['C02', 'City', '5', '7', '+2', '20 min → 14.3 min (28% improvement)'],
        ['E22', 'Intercity', '6', '8', '+2', '20 min → 13.3 min (33% improvement)'],
        ['F25', 'Feeder', '7', '9', '+2', '20 min → 14.7 min (26% improvement)'],
        ['C04', 'City', '7', '4', '−3', '20 min → 35 min (headway increases)'],
        ['F12', 'Feeder', '7', '4', '−3', '20 min → 33 min (headway increases)'],
        ['X66', 'Express', '8', '6', '−2', '20 min → 24 min (headway increases)'],
    ])

add_highlight_box(doc, 'Impact Assessment', [
    '**X28 headway drops 38%** — the single largest service improvement in the proposal.',
    '**Zero capital expenditure** — implementable immediately via operational directive.',
    'Donor routes (C04, F12) have low Pax/km — their riders experience longer waits but from a low base.',
    'This is a **quick win** before structural investments (new buses, bus lanes) can be made.',
])

# --- Headway Modification ---
doc.add_heading('Headway Modification (Time-Variable)', level=2)
add_bold_para(doc,
    'Implement **three-tier timetabling** based on day-of-week demand patterns, achieving '
    '~7-8% effective capacity gain on peak weekdays without additional buses:')

add_table(doc,
    ['Tier', 'Days', 'Demand Index', 'Headway Rule', 'Example (X28)'],
    [
        ['Peak', 'Sunday–Thursday', '1.05–1.10x', 'Base headway', '11–12 min'],
        ['Reduced', 'Friday', '1.04–1.09x', '1.1× base', '13–14 min'],
        ['Minimal', 'Saturday', '0.82–0.88x', '1.3× base', '15–16 min'],
    ])

# --- Corridor Prioritization ---
doc.add_heading('Corridor Prioritization', level=2)
add_table(doc,
    ['Priority', 'Corridor', 'Action', 'Timeline', 'Expected Impact'],
    [
        ['1 (Critical)', 'X28 Express', 'Add 3 buses + initiate bus-lane study', 'Immediate', 'Headway: 20 → 12 min'],
        ['2 (High)', 'Coastal Marina zone', 'Add 2–3 stops + deploy larger buses', '30 days', 'Per-stop load −40%'],
        ['3 (High)', 'X11, X66 Express', 'Headway optimization + contingency', '30 days', 'Better peak coverage'],
        ['4 (Medium)', 'C03, C02 City', 'Reallocate 2 buses each', 'Immediate', 'Headway: 15 → 11 min'],
        ['5 (Monitor)', 'F12, C04 donors', 'Extend headways, monitor impact', 'Ongoing', 'Capacity freed up'],
    ])

# --- Seasonal Rotation ---
doc.add_heading('Seasonal Fleet Rotation', level=2)
add_table(doc,
    ['Season', 'Months', 'Fleet Multiplier', 'Action'],
    [
        ['Winter Peak', 'Nov–Mar', '1.12× baseline', 'Full fleet + 10-15% surge capacity'],
        ['Shoulder', 'Apr–May, Sep–Oct', '1.00× baseline', 'Standard operations'],
        ['Summer', 'Jun–Aug', '0.88× baseline', 'Redeploy 10-15% to scheduled maintenance'],
    ])

add_bold_para(doc,
    'The **22-25% capacity swing** between winter and summer optimizes asset utilization '
    'while maintaining service quality year-round. Summer becomes the **natural maintenance window**.')

doc.add_page_break()

# ============================================================
# APPENDIX
# ============================================================
doc.add_heading('Appendix: Technical Details', level=1)
add_line(doc)

doc.add_heading('Data Pipeline', level=2)
add_table(doc,
    ['Dataset', 'Records', 'Key Fields', 'Role'],
    [
        ['Bus_Routes.csv', '12 routes', 'Type, length, travel time, zone', 'Route metadata'],
        ['Bus_Stops.csv', '56 stops', 'Type, lat/lon, zone', 'Stop metadata'],
        ['Route_Stop_Mapping.csv', '56 mappings', 'Sequence, dwell time', 'Route-stop linkage'],
        ['Train_Ridership.csv', '195,381 rows', 'Daily boarding/alighting', 'Demand data'],
        ['Train_Traffic.csv', '1,277 rows', 'Congestion, speed, incidents', 'Network conditions'],
    ])

add_bold_para(doc,
    'All five datasets were merged into a **single master analytical dataset** (195,381 rows × 27 columns) '
    'with **zero data loss and zero null values**. Feature engineering added: Total_Pax, Year, Month, '
    'DayOfWeek, IsWeekend, Season, and YearMonth.')

doc.add_heading('Analytical Scripts', level=2)
add_table(doc,
    ['Script', 'Purpose', 'Output'],
    [
        ['stage1_pipeline.py', 'Data merge + diagnostics', 'master_analytical_dataset.csv'],
        ['stage1_forecast.py', 'Demand forecasting', 'forecast_h2_2025.csv'],
        ['stage1_corridor_analysis.py', 'Corridor scoring', 'corridor_output.txt'],
        ['stage1_fleet_reallocation.py', 'Fleet optimization', 'fleet_output.txt'],
        ['stage1_visualizations.py', '14 analytical charts', 'charts/ folder'],
        ['growth_decomposition.py', 'Growth breakdown charts', '2 stacked bar charts'],
    ])

# ============================================================
# SAVE
# ============================================================
doc.save(str(OUT_FILE))
print(f"\n{'='*50}")
print(f"Document saved: {OUT_FILE}")
print(f"{'='*50}")
print(f"Size: {OUT_FILE.stat().st_size / 1024:.0f} KB")
